from docx import Document
from docx.shared import Pt

class GenerateDocsReport():

    @classmethod
    def generateReport(cls, report_file, report):
        docx_file = report_file.replace(".html", ".docx")
        doc = Document()
        heading = doc.add_heading(level=0)
        run = heading.add_run(f"Отчёт \n {report}")
        run.font.size = Pt(16)

        doc.save(docx_file)